#!/usr/bin/env python3
"""
generate_report.py — USB File Transfer Activity Driver Report Generator

Uses python-docx to create a professionally formatted Word report.
Install:  pip install python-docx
Output:   scripts/Px_Group10-report-CSC1107.docx
"""

import os
import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = os.path.join(os.path.dirname(__file__),
                           "Px_Group10-report-CSC1107.docx")


def set_cell_shading(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)


def add_styled_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '2E75B6')
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(10)
            if r % 2 == 1:
                set_cell_shading(cell, 'F7F9FC')
    doc.add_paragraph()


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F2F2F2')
    shading.set(qn('w:val'), 'clear')
    pPr.append(shading)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_bullet(doc, text):
    return doc.add_paragraph(text, style='List Bullet')


def add_bold_para(doc, segments):
    p = doc.add_paragraph()
    for is_bold, text in segments:
        run = p.add_run(text)
        if is_bold:
            run.bold = True
    return p


def build_document():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.15

    for level in (1, 2, 3):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Calibri'
        h_style.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79) if level == 1 else RGBColor(0x2E, 0x75, 0xB6)
        h_style.font.bold = True
        sizes = {1: Pt(26), 2: Pt(18), 3: Pt(14)}
        h_style.font.size = sizes[level]
        spacing = {1: (18, 12), 2: (14, 8), 3: (10, 6)}
        h_style.paragraph_format.space_before = Pt(spacing[level][0])
        h_style.paragraph_format.space_after = Pt(spacing[level][1])

    # ═════════════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ═════════════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()
    doc.add_heading('USB File Transfer Activity Driver', level=1)
    doc.add_heading('Design Documentation & Technical Report', level=2)
    doc.add_paragraph()
    add_bold_para(doc, [(True, 'Project: '), (False, 'Project 15 — USB File Transfer Activity Driver')])
    add_bold_para(doc, [(True, 'Course: '), (False, 'CSC1107 — Operating Systems')])
    add_bold_para(doc, [(True, 'Group: '), (False, 'Group 10')])
    add_bold_para(doc, [(True, 'Target Platform: '), (False, 'Raspberry Pi 4, Raspbian 64-bit (Linux 6.x)')])
    add_bold_para(doc, [(True, 'Date: '), (False, datetime.date.today().strftime('%d %B %Y'))])
    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ═════════════════════════════════════════════════════════════════════
    doc.add_heading('Table of Contents', level=1)
    toc = [
        '1. Executive Summary', '2. System Architecture & Design',
        '3. Design Step-by-Step: Building the Driver',
        '4. Kernel Module — Detailed Design',
        '5. User-Space Application — Detailed Design',
        '6. Python Watchdog Monitor', '7. Anomaly Detection Engine',
        '8. Build System & Automation', '9. Communication Protocol (ioctl API)',
        '10. Testing & Verification', '11. Conclusion',
    ]
    for item in toc:
        doc.add_paragraph(item)
    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ═════════════════════════════════════════════════════════════════════
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'This project is a Linux kernel driver that watches file activity on USB storage '
        'devices. Think of it as a security camera for USB drives — it records every file '
        'that is created, changed, or deleted, and raises an alarm if someone copies too '
        'many files too quickly. This is called Data Loss Prevention (DLP), and it helps '
        'stop sensitive information from being stolen via USB sticks.'
    )
    doc.add_paragraph(
        'The system runs on a Raspberry Pi 4 with Raspbian 64-bit and has three parts: '
        '(1) a kernel module (usb_audit.ko) that lives inside the Linux kernel and has '
        'direct access to USB hardware events, (2) a C program (usb_monitor) that shows '
        'statistics and logs in a dashboard, and (3) a Python script (file_tracker.py) '
        'that watches folders for changes and reports them to the kernel module.'
    )
    doc.add_paragraph(
        'The standout feature is automatic mass-copy detection. The kernel module uses '
        'a sliding-window burst counter: if too many file operations happen within a short '
        'time window, it automatically logs a SECURITY ALERT. This works at two levels — '
        'the kernel detects writes autonomously using a kprobe hook on the vfs_write '
        'function, and user-space programs can also inject events for monitoring. Both '
        'levels feed into the same detection engine.'
    )
    doc.add_paragraph(
        'The project has been fully tested on real hardware with a USB drive. An automated '
        'test suite (test_all.sh) verifies all features in one command. All 22 tests pass '
        'consistently.'
    )

    # ═════════════════════════════════════════════════════════════════════
    # 2. SYSTEM ARCHITECTURE
    # ═════════════════════════════════════════════════════════════════════
    doc.add_heading('2. System Architecture & Design', level=1)

    doc.add_heading('2.1 High-Level Architecture', level=2)
    doc.add_paragraph(
        'The system is built in layers. The kernel module (usb_audit.ko) sits at the bottom '
        'with full hardware access. Above it, user-space programs talk to the kernel through '
        'a special file called /dev/usb_audit — just like reading and writing a regular file, '
        'but the data goes to our driver instead of a disk.'
    )
    doc.add_paragraph(
        'There are two ways events reach the kernel module. The first is user-space reporting: '
        'the C monitor or Python script notices a file change and sends a message to '
        '/dev/usb_audit. The second is kernel-level interception: a kprobe (kernel probe) '
        'hooks into the Linux kernel\'s own vfs_write function, so the driver automatically '
        'detects any file write anywhere on the system — no user-space help needed. '
        'Both paths feed into the same log buffer and anomaly detection engine.'
    )

    add_styled_table(doc,
        ['Layer', 'Component', 'Language', 'What It Does'],
        [
            ['User-Space', 'usb_monitor', 'C', 'Dashboard + event injection'],
            ['User-Space', 'file_tracker.py', 'Python 3', 'Real-time folder watching'],
            ['User-Space', 'test_all.sh', 'Bash', 'Automated test suite'],
            ['Kernel-Space', 'usb_audit.ko', 'C (Kernel)', 'Driver + anomaly engine'],
            ['Kernel-Space', 'kprobe hook', 'C (Kernel)', 'Auto-detect file writes'],
        ]
    )

    doc.add_heading('2.2 Data Flow', level=2)
    doc.add_paragraph(
        'When a file is written to a USB drive, two things can happen. First, the kernel\'s '
        'kprobe on vfs_write automatically fires — it checks if the file is on the monitored '
        'USB path, and if so, logs the event with the file size. Second, a user-space program '
        'may also detect the change and send an event string to /dev/usb_audit. In both cases, '
        'the driver records the event in its circular buffer, updates counters, and feeds the '
        'anomaly detector. If too many events happen too fast, a SECURITY ALERT is raised '
        'automatically.'
    )
    doc.add_paragraph(
        'File size tracking works end-to-end: the Python monitor reads the actual file size '
        'and includes it in the event string. The kernel module parses the "(N bytes)" suffix '
        'and stores the size. The kprobe path captures the actual number of bytes written from '
        'the vfs_write return value. This means the statistics always show real byte counts.'
    )

    doc.add_heading('2.3 Concurrency Model', level=2)
    doc.add_paragraph(
        'All shared data is protected by a single mutex lock called audit_mutex. For the '
        'kprobe path, which runs in atomic context where sleeping is not allowed, a special '
        'trylock variant is used — if the lock is already taken, the event is simply skipped '
        'rather than risking a system crash. This is a safe trade-off for a monitoring tool.'
    )

    # ═════════════════════════════════════════════════════════════════════
    # 3. DESIGN STEP-BY-STEP
    # ═════════════════════════════════════════════════════════════════════
    doc.add_heading('3. Design Step-by-Step: Building the Driver', level=1)
    doc.add_paragraph(
        'This section walks through the design process step by step, from the initial '
        'header file all the way to the automated test suite.'
    )

    steps = [
        ('Step 1: Define the Shared Contract (usb_tracker.h)',
         'Before writing any code, we defined the data structures and ioctl commands '
         'that both kernel and user code must agree upon. This header defines event types, '
         'the log entry structure (timestamp, PID, event type, file path, byte size), '
         'statistics counters, and seven ioctl commands.',
         'Design challenge: the GET_LOGS structure is about 35 KB (128 entries x 280 bytes), '
         'exceeding the kernel\'s 14-bit ioctl size limit (max 16,383 bytes). We use a '
         'placeholder type (__u32) in the ioctl macro while the handler still transfers '
         'the full struct.'),
        ('Step 2: Create the Character Device Skeleton',
         'The kernel module starts with alloc_chrdev_region() for a major number, '
         'cdev_init()/cdev_add() to register, and class_create()/device_create() to make '
         '/dev/usb_audit appear. A version check handles the class_create() API change '
         'in kernel 6.4.'),
        ('Step 3: Implement the Circular Log Buffer',
         'A fixed-size array of 128 entries serves as a circular buffer. The head pointer '
         'wraps around; when full, the oldest entry is overwritten. The read operation '
         'properly advances f_pos. Log_count tracks valid entries.'),
        ('Step 4: Add USB Hotplug Monitoring',
         'usb_register_notify() provides callbacks on USB add/remove. Our filter checks '
         'for mass-storage devices (class 0x08) and logs DEVICE_IN / DEVICE_OUT events.'),
        ('Step 5: Add Kernel-Level File Interception with kprobe',
         'A kretprobe on vfs_write intercepts every file write system-wide. The entry '
         'handler saves the file pointer and size; the return handler checks if the file '
         'is on the monitored path and logs the event.',
         'Challenge: kprobes run in atomic context — mutex_lock() would crash. Solved with '
         'audit_log_event_atomic() using mutex_trylock(). If busy, the event is skipped.',
         'We renamed "enable_kprobe" to "enable_vfs_kprobe" to avoid a name clash with '
         'the kernel\'s own enable_kprobe() function.'),
        ('Step 6: Build the ioctl Command Dispatcher',
         'The unlocked_ioctl handler validates the magic number "U" and dispatches seven '
         'commands: GET_STATS, CLEAR_LOGS, SET_PATH, GET_LOGS, RESET_STATS, SET_ANOMALY, '
         'and GET_ANOMALY. RESET_STATS also clears the anomaly ring.'),
        ('Step 7: Implement the Anomaly Detection Engine',
         'A 64-entry ring buffer stores event timestamps. After each CREATE/MODIFY, the '
         'engine counts how many events fall within the sliding window (default 3000 ms). '
         'If the count exceeds the threshold (default 5) and cooldown (5000 ms) has passed, '
         'a SECURITY ALERT is logged. Both logging paths use consistent monotonic timestamps '
         '(ktime_get_ns).'),
        ('Step 8: Create the User-Space Application',
         'usb_monitor (C) connects to /dev/usb_audit and provides interactive and daemon '
         'modes. It sends events with file sizes in the format "C /path (N bytes)" which '
         'the kernel module parses.'),
        ('Step 9: Create the Python Watchdog',
         'file_tracker.py uses the watchdog library for real-time folder monitoring. It '
         'handles create, modify, close, and delete events, sending each to /dev/usb_audit '
         'with actual file sizes. It also has its own anomaly detector.'),
        ('Step 10: Automate Everything',
         'A top-level Makefile builds both components. run.sh handles the full lifecycle. '
         'test_all.sh runs 22 automated tests and produces a pass/fail report.'),
    ]

    for title, *paragraphs in steps:
        doc.add_heading(title, level=2)
        for para in paragraphs:
            doc.add_paragraph(para)

    # ═════════════════════════════════════════════════════════════════════
    # 4. KERNEL MODULE
    # ═════════════════════════════════════════════════════════════════════
    doc.add_heading('4. Kernel Module — Detailed Design', level=1)

    doc.add_heading('4.1 Module Initialisation', level=2)
    doc.add_paragraph(
        'usb_audit_init() performs seven steps: (1) allocate major number, '
        '(2) register cdev, (3) create /dev/usb_audit with kernel 6.4+ version check, '
        '(4) allocate 128-entry log buffer, (5) initialise stats and anomaly state, '
        '(6) register USB notifier, (7) register kretprobe on vfs_write. '
        'The module param "enable_vfs_kprobe" (renamed from "enable_kprobe") allows '
        'disabling the kprobe at load time.'
    )

    doc.add_heading('4.2 Dual Logging Paths', level=2)
    doc.add_paragraph(
        'audit_log_event() uses mutex_lock() for user context (write/ioctl handlers). '
        'audit_log_event_atomic() uses mutex_trylock() for kprobe context where sleeping '
        'is forbidden. Both use ktime_get_ns() for consistent monotonic timestamps. '
        'The atomic path also tracks FILE_DELETE in statistics alongside CREATE and MODIFY.'
    )

    doc.add_heading('4.3 File Operations', level=2)
    add_styled_table(doc,
        ['Operation', 'Handler', 'Description'],
        [
            ['open', 'usb_audit_open', 'Logs the access; always succeeds.'],
            ['release', 'usb_audit_release', 'Logs the close; always succeeds.'],
            ['read', 'usb_audit_read', 'Returns oldest log entry; advances f_pos.'],
            ['write', 'usb_audit_write', 'Parses event code + path + size; logs event.'],
            ['unlocked_ioctl', 'usb_audit_ioctl', 'Dispatches 7 ioctl commands.'],
        ]
    )

    doc.add_heading('4.4 Write Handler & File Size Parsing', level=2)
    doc.add_paragraph(
        'The write handler accepts "<code> <path> (N bytes)". It parses the event code '
        '(C/M/D/A), extracts the path, and uses strrchr() + sscanf() for the optional '
        'size. Events without a size suffix are still logged (size = 0) — a bug that '
        'previously dropped unsized events was fixed.'
    )

    doc.add_heading('4.5 USB Notifier & Error Handling', level=2)
    doc.add_paragraph(
        'The notifier filters for mass-storage devices (class 0x08 and per-interface 0x00) '
        'and logs vendor/product IDs. Error handling follows Linux conventions with negative '
        'errno returns and goto-based init unwinding. printk() at KERN_INFO/WARNING/ERR '
        'levels provides dmesg visibility.'
    )

    # ═════════════════════════════════════════════════════════════════════
    # 5. USER-SPACE APPLICATION
    # ═════════════════════════════════════════════════════════════════════
    doc.add_heading('5. User-Space Application — Detailed Design', level=1)

    doc.add_heading('5.1 Overview', level=2)
    doc.add_paragraph(
        'usb_monitor.c (~530 lines) communicates with the driver via open(), write(), '
        'and ioctl() on /dev/usb_audit. It provides an interactive test menu (11 commands: '
        'C/M/D for events, A for alert, T for threshold config, S/L for stats/logs, '
        'R/X for reset/clear, Q for quit) and a daemon dashboard mode (-d flag, 2-second '
        'refresh). Command-line flags set path (-p), threshold (-t), and window (-w). '
        'A null-termination typo in the mount path was fixed.'
    )

    doc.add_heading('5.2 Event Formatting', level=2)
    doc.add_paragraph(
        'send_event() formats events as "<code> <path> (<size> bytes)\\n" and writes to '
        '/dev/usb_audit. The kernel parses this format. Interactive mode uses a placeholder '
        'size (1024 bytes); the Python monitor sends actual sizes.'
    )

    # ═════════════════════════════════════════════════════════════════════
    # 6. PYTHON WATCHDOG
    # ═════════════════════════════════════════════════════════════════════
    doc.add_heading('6. Python Watchdog Monitor', level=1)
    doc.add_paragraph(
        'file_tracker.py uses the watchdog library for real-time folder monitoring. '
        'It runs in the background and automatically reports file events to the kernel '
        'module through /dev/usb_audit.'
    )

    doc.add_heading('6.1 Event Handlers', level=2)
    doc.add_paragraph(
        'Four handlers: on_created() sends "C" with real file size; on_modified() tracks '
        'growing sizes; on_closed() sends "M" with final size; on_deleted() sends "D" '
        'with size 0. All feed the local anomaly detector.'
    )

    doc.add_heading('6.2 Kernel Integration & Local Detection', level=2)
    doc.add_paragraph(
        'send_kernel_event() writes to /dev/usb_audit, catching errors if the module '
        'is not loaded. The script works with or without the kernel module. The local '
        'anomaly detector mirrors the kernel logic: sliding window (default 3.0s) with '
        'a threshold of 5 operations.'
    )

    # ═════════════════════════════════════════════════════════════════════
    # 7. ANOMALY DETECTION
    # ═════════════════════════════════════════════════════════════════════
    doc.add_heading('7. Anomaly Detection Engine', level=1)

    doc.add_heading('7.1 How It Works', level=2)
    doc.add_paragraph(
        'The engine works like a speed camera: count file operations within a time window, '
        'raise an alarm if too many. A 64-entry ring buffer stores timestamps. After each '
        'CREATE/MODIFY, it counts how many entries fall within the sliding window '
        '(default 3000 ms). If the count exceeds the threshold (default 5) and the cooldown '
        '(5000 ms) has passed, a SECURITY ALERT is logged. The cooldown prevents spam.'
    )

    doc.add_heading('7.2 Configuration & Verification', level=2)
    doc.add_paragraph(
        'Fully configurable at runtime via T command, flags, or SET_ANOMALY ioctl. '
        'GET_ANOMALY returns burst count and alert status. Verified: threshold=2, '
        'window=5000ms — 3 rapid events trigger the alert. Both the monitor displays '
        'a red SECURITY ALERT banner and dmesg confirms with burst details.'
    )

    doc.add_heading('7.3 Three Detection Layers', level=2)
    doc.add_paragraph(
        '(1) Kernel-side anomaly_check_burst() from both logging paths. '
        '(2) User-space polling via GET_ANOMALY ioctl. '
        '(3) Python-side sliding window in file_tracker.py. '
        'If any layer detects a burst, the alert is visible.'
    )

    # ═════════════════════════════════════════════════════════════════════
    # 8. BUILD SYSTEM
    # ═════════════════════════════════════════════════════════════════════
    doc.add_heading('8. Build System & Automation', level=1)

    doc.add_heading('8.1 Makefile', level=2)
    doc.add_paragraph(
        'Top-level targets: all (kernel + user), kernel (Kbuild via src_kernel/), '
        'user (gcc -Wall -Wextra -O2), clean. The kernel Makefile uses '
        'obj-m := usb_audit.o and auto-locates headers.'
    )

    doc.add_heading('8.2 Deployment Script (run.sh)', level=2)
    doc.add_paragraph(
        'Full lifecycle in one command: checks root and kernel headers, builds both '
        'components, loads the module, verifies /dev/usb_audit, shows kernel logs, '
        'launches monitor, auto-cleans on Ctrl+C. Colour-coded output.'
    )

    doc.add_heading('8.3 Automated Test Suite (test_all.sh)', level=2)
    doc.add_paragraph(
        '22 tests across 10 categories: module loading, device node, kernel log, kprobe, '
        'notifier, monitor connection, event injection (C/M/D), statistics accuracy, log '
        'buffer, anomaly detection, USB drive I/O, clean shutdown. Handles edge cases and '
        'resets stats between tests.'
    )
    add_code_block(doc, 'sudo ./scripts/test_all.sh   # Run all 22 tests')

    # ═════════════════════════════════════════════════════════════════════
    # 9. IOCTL API
    # ═════════════════════════════════════════════════════════════════════
    doc.add_heading('9. Communication Protocol (ioctl API)', level=1)
    doc.add_paragraph(
        'All kernel/user communication uses ioctl() on /dev/usb_audit with magic number '
        '"U" (0x55). Commands are defined in usb_tracker.h.'
    )

    add_styled_table(doc,
        ['Command', 'Direction', 'Data', 'Purpose'],
        [
            ['GET_STATS (1)', 'Kernel to User', 'Stats struct', 'Read all counters'],
            ['CLEAR_LOGS (2)', 'None', 'None', 'Empty log buffer'],
            ['SET_PATH (3)', 'User to Kernel', 'String (256)', 'Set USB mount path'],
            ['GET_LOGS (4)', 'Both ways', 'Log entries', 'Read event history'],
            ['RESET_STATS (5)', 'None', 'None', 'Zero counters + anomaly ring'],
            ['SET_ANOMALY (6)', 'User to Kernel', 'Config struct', 'Set threshold + window'],
            ['GET_ANOMALY (7)', 'Kernel to User', 'Status struct', 'Check burst + alert'],
        ]
    )
    doc.add_paragraph(
        'GET_LOGS struct is 35 KB, exceeding the 14-bit ioctl size limit (16,383 bytes). '
        'We use __u32 placeholder in the macro; actual data transfer uses the full struct.'
    )

    # ═════════════════════════════════════════════════════════════════════
    # 10. TESTING
    # ═════════════════════════════════════════════════════════════════════
    doc.add_heading('10. Testing & Verification', level=1)

    doc.add_heading('10.1 Automated Suite', level=2)
    doc.add_paragraph(
        'sudo ./scripts/test_all.sh produces a pass/fail report. '
        'On Raspberry Pi 4 with real USB drive, all 22 tests pass.'
    )

    doc.add_heading('10.2 Module Lifecycle', level=2)
    for cmd in [
        'sudo insmod src_kernel/usb_audit.ko   # Load the driver',
        'lsmod | grep usb_audit                # Confirm loaded',
        'ls -la /dev/usb_audit                 # Check device node',
        'dmesg | tail -20                      # Kernel messages',
        'sudo rmmod usb_audit                  # Unload',
    ]:
        add_code_block(doc, cmd)

    doc.add_heading('10.3 Event Injection', level=2)
    for cmd in [
        'C /media/colossalblue/RAZER/doc.pdf   # Log file CREATE',
        'M /media/colossalblue/RAZER/data.txt  # Log file MODIFY',
        'D /media/colossalblue/RAZER/old.txt   # Log file DELETE',
        'S                                       # Show statistics',
        'L                                       # View logs',
    ]:
        add_code_block(doc, cmd)

    doc.add_heading('10.4 Mass-Copy Alert', level=2)
    for cmd in [
        'T 2 5000                    # Threshold=2, window=5000ms',
        'C /media/usb/a.txt          # Event 1',
        'C /media/usb/b.txt          # Event 2',
        'C /media/usb/c.txt          # Event 3 -- ALERT!',
    ]:
        add_code_block(doc, cmd)
    doc.add_paragraph(
        'Monitor shows red SECURITY ALERT. dmesg: "[usb_audit] *** SECURITY ALERT *** '
        'Mass-copy detected! 3 file ops within 5000 ms (threshold=2)".'
    )

    doc.add_heading('10.5 Real USB Drive', level=2)
    doc.add_paragraph(
        'Real file operations (echo, cp, rm) on a mounted USB drive generated events '
        'in both the C dashboard and dmesg. kprobe auto-detected writes. Hotplug events '
        'logged as DEVICE_IN/DEVICE_OUT.'
    )

    doc.add_heading('10.6 Verification Checklist', level=2)
    add_styled_table(doc,
        ['#', 'Feature', 'How to Verify', 'Status'],
        [
            ['1', 'Compiles', 'make all exits 0', 'Passed'],
            ['2', 'Loads', 'insmod; /dev/usb_audit created', 'Passed'],
            ['3', 'kprobe active', 'dmesg shows kretprobe registered', 'Passed'],
            ['4', 'Monitor connects', 'usb_monitor shows Connected', 'Passed'],
            ['5', 'Events inject', 'C/M/D produce Sent event', 'Passed'],
            ['6', 'Stats accurate', 'S shows correct counters', 'Passed'],
            ['7', 'Logs work', 'L shows timestamped entries', 'Passed'],
            ['8', 'Anomaly alert', '3 rapid events trigger alert', 'Passed'],
            ['9', 'Kernel logs alert', 'dmesg shows alert details', 'Passed'],
            ['10', 'USB drive I/O', 'Write/read real files on USB', 'Passed'],
            ['11', 'Clean unload', 'rmmod; /dev/usb_audit removed', 'Passed'],
            ['12', '22 auto-tests', 'test_all.sh 22/22 passed', 'Passed'],
        ]
    )

    # ═════════════════════════════════════════════════════════════════════
    # 11. CONCLUSION
    # ═════════════════════════════════════════════════════════════════════
    doc.add_heading('11. Conclusion', level=1)
    doc.add_paragraph(
        'The USB File Transfer Activity Driver meets all project requirements with '
        'advanced features. The kernel module monitors USB storage at device (hotplug) '
        'and file (create/modify/delete) levels. User-space tools provide interactive '
        'and automated interfaces.'
    )
    doc.add_paragraph('Key achievements:')
    for achievement in [
        'kretprobe on vfs_write for autonomous kernel-level file interception.',
        'Dual-path logging: atomic-safe (kprobe) + regular (user context).',
        'Configurable sliding-window anomaly detector with cooldown.',
        'End-to-end file size tracking with real byte counts.',
        'Full ioctl API: 7 commands for stats, logs, path, and anomaly config.',
        'Linux 6.4+ compatibility via preprocessor version check.',
        '128-entry circular log buffer with proper file position handling.',
        'Three-layer anomaly detection: kernel + user-space + Python.',
        'Automated test suite: 22/22 tests pass on real hardware.',
        'One-command deployment: make all && sudo ./scripts/run.sh.',
    ]:
        add_bullet(doc, achievement)

    doc.add_paragraph()
    doc.add_paragraph(
        'Tested on Raspberry Pi 4 with kernel 6.18.26-v71-CSC1107_CUSTOM_KERNEL+ '
        'and a real USB drive. All 22 automated tests pass. The system detects USB '
        'hotplug events, logs file operations with accurate timestamps and sizes, '
        'and raises security alerts on mass-copy behaviour.'
    )
    doc.add_paragraph(
        'This project demonstrates kernel module development, character device drivers, '
        'kernel/user ioctl communication, mutex-based concurrency, atomic context '
        'programming, kprobe function hooking, USB subsystem integration, and automated '
        'build/test pipelines.'
    )

    return doc


def main():
    doc = build_document()
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    size = os.path.getsize(OUTPUT_PATH)
    print(f'Report generated: {OUTPUT_PATH}')
    print(f'Size: {size:,} bytes')


if __name__ == '__main__':
    main()
